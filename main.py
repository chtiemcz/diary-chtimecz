import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageTk
from datetime import datetime

image_path = "C:/Users/xtime/Desktop/a100do/dnewnik/bg.jpg"  # Переменная для хранения пути к изображению фона

def save_to_docx(entries):
    doc = Document()

    for entry in entries:
        doc.add_paragraph(entry)

    if image_path:
        doc.add_picture(image_path, width=Inches(6), height=Inches(4))

    today = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")  # Текущая дата и время
    file_path = f"Дневник_{today}.docx"
    doc.save(file_path)
    print(f"Дневник сохранен в файл: {file_path}")

def browse_image():
    global image_path
    file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg")])
    if file_path:
        image_path = file_path
        # Установка фонового изображения
        img = Image.open(image_path)
        img = img.resize((root.winfo_width(), root.winfo_height()))
        img = ImageTk.PhotoImage(img)
        background_label.configure(image=img)
        background_label.image = img

# Создание окна
root = tk.Tk()
root.title("Дневник")

# Установка размера окна
root.geometry("750x410")  # Установите размер окна по вашему усмотрению
root.resizable(False, False)
# Создание метки для фонового изображения
background_label = tk.Label(root)
background_label.place(x=0, y=0, relwidth=1, relheight=1)

# Установка фонового изображения
img = Image.open(image_path)
img = img.resize((root.winfo_width(), root.winfo_height()))
img = ImageTk.PhotoImage(img)
background_label.configure(image=img)
background_label.image = img

# Создание текстового поля для записей
text_entry = tk.Text(root, height=15, width=103, font=("Comic Sans MS", 8))  # Задание шрифта текста
text_entry.place(x=10, y=25)

# Кнопка "Сохранить"
save_button = tk.Button(root, text="Сохранить", command=lambda: save_to_docx(text_entry.get("1.0", "end-1c").split("\n")))
save_button.place(x=340, y=310)

# Кнопка "Загрузить картинку"
load_image_button = tk.Button(root, text="Загрузить картинку", command=browse_image)
load_image_button.place(relx=0.5, rely=0.7, anchor="center")

root.mainloop()